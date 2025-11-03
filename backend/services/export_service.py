"""
Export Service for generating DOCX, XLSX, and PDF files
"""

import os
import tempfile
from typing import Dict
from datetime import datetime

# Import libraries with error handling
try:
    from docx import Document
    from docx.shared import Pt, Inches, RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    from openpyxl import Workbook
    from openpyxl.styles import Font, Alignment, PatternFill
    XLSX_AVAILABLE = True
except ImportError:
    XLSX_AVAILABLE = False

try:
    from reportlab.lib.pagesizes import letter, A4
    from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
    from reportlab.lib.units import inch
    from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
    from reportlab.lib import colors
    from reportlab.pdfgen import canvas
    from PyPDF2 import PdfReader, PdfWriter
    import io
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False


class ExportService:
    """Service for exporting data to various formats"""
    
    def __init__(self):
        self.temp_dir = tempfile.gettempdir()
        self.docx_available = DOCX_AVAILABLE
        self.xlsx_available = XLSX_AVAILABLE
        self.pdf_available = PDF_AVAILABLE
    
    def is_available(self) -> bool:
        """Check if export services are available"""
        return self.docx_available or self.xlsx_available or self.pdf_available
    
    async def export_to_docx(
        self,
        extracted_data: Dict,
        translations: Dict,
        steel_equivalents: Dict = None
    ) -> str:
        """Export data to DOCX format"""
        if not self.docx_available:
            raise ImportError("python-docx not installed. Install with: pip install python-docx")
        
        if steel_equivalents is None:
            steel_equivalents = {}
        
        # Create document
        doc = Document()
        
        # Title
        title = doc.add_heading('Drawing Analysis Report', 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Date
        doc.add_paragraph(f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}")
        doc.add_paragraph("")
        
        # Materials section
        doc.add_heading('Materials', level=1)
        if extracted_data.get("materials"):
            for i, material in enumerate(extracted_data["materials"]):
                p = doc.add_paragraph(f"• {material}", style='List Bullet')
                if i < len(translations.get("materials", [])):
                    translated = translations["materials"][i]
                    if translated != material:
                        p.add_run(f" → {translated}").italic = True
                
                # Add steel equivalents if available
                if material in steel_equivalents:
                    equiv = steel_equivalents[material]
                    doc.add_paragraph(
                        f"  ASTM: {equiv.get('astm', 'N/A')}, "
                        f"ISO: {equiv.get('iso', 'N/A')}, "
                        f"GB/T: {equiv.get('gbt', 'N/A')}",
                        style='List Bullet 2'
                    )
        
        # GOST/OST/TU section
        if extracted_data.get("standards"):
            doc.add_heading('Standards (GOST/OST/TU)', level=1)
            for standard in extracted_data["standards"]:
                doc.add_paragraph(f"• {standard}", style='List Bullet')
        
        # Surface roughness (Ra)
        if extracted_data.get("ra"):
            doc.add_heading('Surface Roughness (Ra)', level=1)
            doc.add_paragraph(f"Ra: {', '.join(extracted_data['ra'])}")
        
        # Fits
        if extracted_data.get("fits"):
            doc.add_heading('Fits', level=1)
            doc.add_paragraph(f"Fits: {', '.join(extracted_data['fits'])}")
        
        # Heat treatment
        if extracted_data.get("heatTreatment"):
            doc.add_heading('Heat Treatment', level=1)
            doc.add_paragraph(f"Heat Treatment: {extracted_data['heatTreatment']}")
        
        # Save to temporary file
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"drawing_analysis_{timestamp}.docx"
        filepath = os.path.join(self.temp_dir, filename)
        doc.save(filepath)
        
        return filepath
    
    async def export_to_xlsx(
        self,
        extracted_data: Dict,
        translations: Dict,
        steel_equivalents: Dict = None
    ) -> str:
        """Export data to XLSX format"""
        if not self.xlsx_available:
            raise ImportError("openpyxl not installed. Install with: pip install openpyxl")
        
        if steel_equivalents is None:
            steel_equivalents = {}
        
        # Create workbook
        wb = Workbook()
        ws = wb.active
        ws.title = "Drawing Analysis"
        
        # Header style
        header_fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
        header_font = Font(bold=True, color="FFFFFF", size=12)
        
        # Title
        ws.merge_cells('A1:D1')
        title_cell = ws['A1']
        title_cell.value = "Drawing Analysis Report"
        title_cell.font = Font(bold=True, size=14)
        title_cell.alignment = Alignment(horizontal='center')
        
        ws['A2'] = f"Generated: {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
        
        row = 4
        
        # Materials section
        ws.merge_cells(f'A{row}:D{row}')
        ws[f'A{row}'] = "Materials"
        ws[f'A{row}'].font = header_font
        ws[f'A{row}'].fill = header_fill
        row += 1
        
        ws['A5'] = "Original"
        ws['B5'] = "Translated"
        ws['C5'] = "ASTM"
        ws['D5'] = "ISO"
        ws['E5'] = "GB/T"
        for col in ['A', 'B', 'C', 'D', 'E']:
            ws[f'{col}5'].font = Font(bold=True)
            ws[f'{col}5'].fill = header_fill
        
        row = 6
        if extracted_data.get("materials"):
            for i, material in enumerate(extracted_data["materials"]):
                ws[f'A{row}'] = material
                if i < len(translations.get("materials", [])):
                    ws[f'B{row}'] = translations["materials"][i]
                
                if material in steel_equivalents:
                    equiv = steel_equivalents[material]
                    ws[f'C{row}'] = equiv.get('astm', '')
                    ws[f'D{row}'] = equiv.get('iso', '')
                    ws[f'E{row}'] = equiv.get('gbt', '')
                
                row += 1
        
        # Standards
        row += 2
        ws.merge_cells(f'A{row}:D{row}')
        ws[f'A{row}'] = "Standards (GOST/OST/TU)"
        ws[f'A{row}'].font = header_font
        ws[f'A{row}'].fill = header_fill
        row += 1
        
        if extracted_data.get("standards"):
            for standard in extracted_data["standards"]:
                ws[f'A{row}'] = standard
                row += 1
        
        # Surface roughness
        if extracted_data.get("ra"):
            row += 1
            ws.merge_cells(f'A{row}:D{row}')
            ws[f'A{row}'] = "Surface Roughness (Ra)"
            ws[f'A{row}'].font = header_font
            ws[f'A{row}'].fill = header_fill
            row += 1
            ws[f'A{row}'] = ', '.join(extracted_data['ra'])
        
        # Fits
        if extracted_data.get("fits"):
            row += 2
            ws.merge_cells(f'A{row}:D{row}')
            ws[f'A{row}'] = "Fits"
            ws[f'A{row}'].font = header_font
            ws[f'A{row}'].fill = header_fill
            row += 1
            ws[f'A{row}'] = ', '.join(extracted_data['fits'])
        
        # Heat treatment
        if extracted_data.get("heatTreatment"):
            row += 2
            ws.merge_cells(f'A{row}:D{row}')
            ws[f'A{row}'] = "Heat Treatment"
            ws[f'A{row}'].font = header_font
            ws[f'A{row}'].fill = header_fill
            row += 1
            ws[f'A{row}'] = extracted_data['heatTreatment']
        
        # Auto-adjust column widths
        for column in ws.columns:
            max_length = 0
            column_letter = column[0].column_letter
            for cell in column:
                try:
                    if len(str(cell.value)) > max_length:
                        max_length = len(str(cell.value))
                except:
                    pass
            adjusted_width = min(max_length + 2, 50)
            ws.column_dimensions[column_letter].width = adjusted_width
        
        # Save to temporary file
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        filename = f"drawing_analysis_{timestamp}.xlsx"
        filepath = os.path.join(self.temp_dir, filename)
        wb.save(filepath)
        
        return filepath
    
    async def export_to_pdf(
        self,
        pdf_content: bytes,
        extracted_data: Dict,
        translations: Dict,
        steel_equivalents: Dict = None
    ) -> str:
        """Export PDF with English overlay"""
        if not self.pdf_available:
            raise ImportError("reportlab and PyPDF2 not installed. Install with: pip install reportlab PyPDF2")
        
        if steel_equivalents is None:
            steel_equivalents = {}
        
        # Save original PDF temporarily
        timestamp = datetime.now().strftime("%Y%m%d_%H%M%S")
        temp_pdf = os.path.join(self.temp_dir, f"original_{timestamp}.pdf")
        with open(temp_pdf, "wb") as f:
            f.write(pdf_content)
        
        # Read original PDF
        reader = PdfReader(temp_pdf)
        writer = PdfWriter()
        
        # Create overlay PDF with annotations
        overlay_pdf = os.path.join(self.temp_dir, f"overlay_{timestamp}.pdf")
        c = canvas.Canvas(overlay_pdf, pagesize=A4)
        
        # Add translation overlay (simplified - in production you'd want more sophisticated positioning)
        y_position = 750
        c.setFont("Helvetica", 10)
        c.setFillColor(colors.red)
        
        # Add materials translations
        if extracted_data.get("materials") and translations.get("materials"):
            c.drawString(50, y_position, "Materials (English):")
            y_position -= 20
            for i, material in enumerate(extracted_data["materials"]):
                if i < len(translations["materials"]):
                    translated = translations["materials"][i]
                    if translated != material:
                        c.drawString(70, y_position, f"{material} → {translated}")
                        y_position -= 15
        
        # Add other translations
        if extracted_data.get("standards"):
            y_position -= 10
            c.drawString(50, y_position, "Standards:")
            y_position -= 20
            for standard in extracted_data["standards"]:
                c.drawString(70, y_position, standard)
                y_position -= 15
        
        c.save()
        
        # Merge overlay with original
        overlay_reader = PdfReader(overlay_pdf)
        
        for page_num, page in enumerate(reader.pages):
            if page_num < len(overlay_reader.pages):
                page.merge_page(overlay_reader.pages[page_num])
            writer.add_page(page)
        
        # Save final PDF
        output_pdf = os.path.join(self.temp_dir, f"drawing_analysis_{timestamp}.pdf")
        with open(output_pdf, "wb") as f:
            writer.write(f)
        
        # Cleanup temp files
        try:
            os.remove(temp_pdf)
            os.remove(overlay_pdf)
        except:
            pass
        
        return output_pdf

